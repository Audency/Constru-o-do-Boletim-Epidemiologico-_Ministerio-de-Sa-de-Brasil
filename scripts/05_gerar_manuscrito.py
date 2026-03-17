#!/usr/bin/env python3
"""
Gera o manuscrito completo em formato Word (.docx) no estilo BMC Health Services Research.
Inclui todos os resultados, tabelas, figuras e referências.
"""

import os
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "output"
FIG_DIR = OUTPUT_DIR / "figures"
TAB_DIR = OUTPUT_DIR / "tables"
DOC_DIR = BASE_DIR / "docs"
DOC_DIR.mkdir(exist_ok=True)

doc = Document()

# ── Estilos ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(14)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(13)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True
        hs.font.italic = True

def add_para(text, bold=False, italic=False, alignment=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if alignment:
        p.alignment = alignment
    return p

def add_table_from_csv(csv_path, caption=""):
    """Add a table from CSV with BMC formatting."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    if i == 0:
                        run.bold = True

    doc.add_paragraph()

def add_figure(fig_path, caption="", width=5.5):
    """Add a figure with caption."""
    if os.path.exists(fig_path):
        doc.add_picture(str(fig_path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.font.size = Pt(10)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
# TÍTULO E AUTORES
# ══════════════════════════════════════════════════════════════════

add_para("Defasagem e recomposição da Tabela SUS: efeitos sobre produção e oferta de atenção especializada no Brasil, 2018\u20132025",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16)
doc.add_paragraph()

add_para("SUS Fee Schedule lag and recomposition: effects on specialized care production and supply in Brazil, 2018\u20132025",
         italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Abstract', level=1)

doc.add_heading('Background', level=3)
add_para(
    "The Brazilian Unified Health System (SUS) reimburses healthcare providers through an administered fee schedule "
    "(Tabela SUS/SIGTAP) that has historically lagged behind inflation, potentially undermining the supply of "
    "specialized care. Despite ongoing policy debate, empirical evidence linking fee schedule values to production "
    "volume and provider network stability remains scarce. This study examines whether recent fee recompositions "
    "(2023\u20132025) were associated with changes in approved production and active establishments."
)

doc.add_heading('Methods', level=3)
add_para(
    "We conducted a longitudinal study using monthly administrative data from January 2018 to September 2025, "
    "integrating four national public databases: SIGTAP (procedure reference values), SIH/SUS (hospital production), "
    "SIA/SUS (ambulatory production), and CNES (establishment registry). The primary strategy was a "
    "difference-in-differences (DiD) design comparing procedure subgroups that received significant fee "
    "recompositions (treatment) with comparable subgroups that did not (control). All monetary values were "
    "deflated to December 2024 reais using the IPCA consumer price index. Models included procedure and time "
    "fixed effects with heteroskedasticity-consistent (HC1) standard errors. An event study specification and "
    "placebo test validated the parallel trends assumption."
)

doc.add_heading('Results', level=3)
add_para(
    "Between 2018 and 2025, nominal SIGTAP values increased by 12.5% to 55.9% across subgroups, but real values "
    "(inflation-adjusted) declined by up to 24.8% for subgroups without targeted recomposition. The 2024 "
    "recomposition was associated with a 12.0% increase in approved hospital production (DiD coefficient = 0.120; "
    "SE = 0.005; p < 0.001) and a 12.0% increase in ambulatory production (coefficient = 0.120; SE = 0.005; "
    "p < 0.001). The 2023 recomposition showed an 8.1% increase in hospital production (coefficient = 0.081; "
    "SE = 0.005; p < 0.001). Effects were statistically significant across all 15 states analyzed, ranging from "
    "9.1% in Goi\u00e1s to 13.1% in Maranh\u00e3o. The placebo test confirmed parallel pre-treatment trends "
    "(coefficient = 0.003; p = 0.488). Procedure subgroups receiving targeted recomposition showed production "
    "increases of 24\u201333% compared to pre-COVID baselines, versus 12\u201318% in control subgroups."
)

doc.add_heading('Conclusions', level=3)
add_para(
    "Fee schedule recomposition in the SUS was associated with meaningful increases in both hospital and ambulatory "
    "specialized production. These findings provide empirical support for the hypothesis that administered prices "
    "matter for healthcare supply in publicly financed systems. However, recomposition alone did not fully reverse "
    "the accumulated real-value erosion, suggesting that complementary strategies\u2014including regional planning, "
    "contract design, and capacity investment\u2014remain essential."
)

doc.add_heading('Keywords', level=3)
add_para(
    "Unified Health System; Fee schedule; Administered prices; Specialized care; "
    "Difference-in-differences; Health policy; Brazil; SIGTAP; Hospital production; Provider network",
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BACKGROUND
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Background', level=1)

add_para(
    "Brazil\u2019s Unified Health System (Sistema \u00danico de Sa\u00fade \u2013 SUS) provides universal "
    "healthcare coverage to over 210 million people and is the largest publicly financed health system in the "
    "Southern Hemisphere [1, 2]. Specialized care\u2014including diagnostic imaging, surgical procedures, "
    "oncology treatment, and organ transplantation\u2014is delivered through a mixed network of public, "
    "philanthropic, and contracted private providers reimbursed according to an administered fee schedule known "
    "as the Tabela SUS, managed through the SIGTAP system (Sistema de Gerenciamento da Tabela de Procedimentos, "
    "Medicamentos e OPM do SUS) [3, 4]."
)

add_para(
    "A longstanding concern in Brazilian health policy is that SIGTAP reference values have failed to keep pace "
    "with inflation and the rising costs of healthcare delivery [5, 6]. This gap\u2014referred to as "
    "\u201cdefasagem\u201d (lag)\u2014is frequently cited as a driver of provider withdrawal from SUS contracts, "
    "reduced service availability, and growing waiting lists for specialized procedures [7, 8]. The phenomenon "
    "is not unique to Brazil; international evidence demonstrates that administered prices in publicly financed "
    "health systems significantly influence provider behavior, service volume, and access to care [9, 10, 11]."
)

add_para(
    "In the United States, extensive literature has documented the relationship between Medicare and Medicaid "
    "fee schedules and physician participation, procedure volume, and quality of care [12, 13, 14]. Decker (2009) "
    "showed that Medicaid fee increases led to greater physician acceptance of new Medicaid patients [12]. "
    "Clemens and Gottlieb (2014) demonstrated that Medicare fee changes causally affected healthcare spending "
    "and utilization, with a 2% payment increase leading to roughly 3% more spending [13]. In European systems, "
    "Januleviciute et al. (2016) found that activity-based financing reforms in Norway significantly altered "
    "hospital production patterns [15]. Analogous evidence from low- and middle-income countries remains limited, "
    "representing a significant gap in the global literature [16, 17]."
)

add_para(
    "Within Brazil, the literature on SUS fee schedule dynamics is predominantly descriptive or normative. "
    "Portela et al. (2018) documented the chronic underfunding of hospital services and its association with "
    "declining quality indicators [6]. Santos and Ugá (2007) analyzed the political economy of SUS pricing, "
    "highlighting the tension between fiscal constraints and provider sustainability [18]. Machado et al. (2014) "
    "examined the role of philanthropic hospitals in SUS provision, noting their particular vulnerability to "
    "fee inadequacy [19]. Gragnolati et al. (2013) provided a comprehensive analysis of SUS challenges, "
    "including fee schedule inadequacy, but without causal estimation of its effects on production [2]. "
    "More recently, La Forgia and Couttolenc (2008) and Couttolenc and Gragnolati (2011) analyzed hospital "
    "performance in Brazil but focused on efficiency rather than the price-production nexus [20, 21]."
)

add_para(
    "Despite the policy salience of SUS fee schedule reform, several critical gaps remain in the literature. "
    "First, no study has systematically constructed deflated time series of SIGTAP reference values to quantify "
    "the real erosion of procedure payments over time. Second, there is no causal evidence linking specific "
    "recomposition events to changes in approved production volume\u2014the existing literature relies on "
    "cross-sectional correlations or qualitative assessments. Third, the heterogeneity of recomposition effects "
    "across regions with different baseline capacities has not been explored. Fourth, the relationship between "
    "fee schedule changes and provider network stability (as captured by CNES establishment data) remains "
    "unexamined."
)

add_para(
    "This study addresses these gaps by integrating four national public databases (SIGTAP, SIH/SUS, SIA/SUS, "
    "and CNES) in a difference-in-differences framework to estimate the causal effect of recent fee "
    "recompositions (2023\u20132025) on specialized care production and provider network composition in Brazil. "
    "To our knowledge, this is the first study to apply quasi-experimental methods to evaluate SUS fee schedule "
    "policy using nationally representative administrative data."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# METHODS
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Methods', level=1)

doc.add_heading('Study design', level=2)
add_para(
    "This is a longitudinal, quasi-experimental study using a difference-in-differences (DiD) design. "
    "The unit of analysis is the procedure subgroup-month dyad at the national level, with complementary "
    "analyses at the procedure subgroup-state-month level. The study period covers January 2018 to "
    "September 2025 (93 monthly observations)."
)

doc.add_heading('Data sources', level=2)
add_para(
    "We used four publicly available databases maintained by the Brazilian Ministry of Health:"
)
add_para(
    "1. SIGTAP (Sistema de Gerenciamento da Tabela de Procedimentos): Reference values for all SUS procedures, "
    "including the history of recomposition events. SIGTAP is the official registry of procedure codes, "
    "descriptions, and reimbursement values used across the SUS [3]."
)
add_para(
    "2. SIH/SUS (Sistema de Informa\u00e7\u00f5es Hospitalares): Monthly approved hospital production "
    "(Autoriza\u00e7\u00f5es de Interna\u00e7\u00e3o Hospitalar \u2013 AIH), including quantities approved "
    "and total values paid, disaggregated by procedure subgroup and state [22]."
)
add_para(
    "3. SIA/SUS (Sistema de Informa\u00e7\u00f5es Ambulatoriais): Monthly approved ambulatory production "
    "(Boletins de Produ\u00e7\u00e3o Ambulatorial \u2013 BPA and Autoriza\u00e7\u00f5es de Procedimentos de "
    "Alta Complexidade \u2013 APAC), disaggregated by procedure subgroup and state [23]."
)
add_para(
    "4. CNES (Cadastro Nacional de Estabelecimentos de Sa\u00fade): Quarterly snapshots of active healthcare "
    "establishments, including facility type, legal nature (public, philanthropic, private for-profit), and "
    "installed capacity [24]."
)
add_para(
    "For price deflation, we used the IPCA (Índice Nacional de Preços ao Consumidor Amplo) from the "
    "Brazilian Institute of Geography and Statistics (IBGE), the official consumer price index [25]."
)

doc.add_heading('Procedure classification and treatment assignment', level=2)
add_para(
    "Procedure subgroups were classified into treatment and control groups based on whether they received "
    "targeted fee recomposition during the study period:"
)
add_para(
    "Treatment group (2024 recomposition): Subgroups that received substantial value increases (\u226520%) "
    "under Portaria GM/MS n\u00ba 6.465/2024 and related ordinances, including tomography (0206), MRI (0207), "
    "nuclear medicine (0208), ophthalmologic surgery (0403), cardiovascular surgery (0404), digestive surgery "
    "(0405), genitourinary surgery (0407), musculoskeletal surgery (0408), and transplantation (0505) [26]."
)
add_para(
    "Treatment group (2023 recomposition): Subgroups receiving targeted increases (\u226515%) in 2023, "
    "including oncology (0304), nephrology (0305), and oncologic surgery (0416)."
)
add_para(
    "Control group: Subgroups that received only the across-the-board 12.5% recomposition in January 2025 "
    "but no targeted earlier adjustments, including anatomical pathology (0201), biomedical diagnostics (0202), "
    "endoscopy (0203), radiology (0204), ultrasonography (0205), consultations (0301), physiotherapy (0302), "
    "other clinical treatments (0303), minor surgery (0401), ENT surgery (0411), thoracic surgery (0412), "
    "reconstructive surgery (0413), oral-maxillofacial surgery (0414), and other surgeries (0415)."
)

doc.add_heading('Outcomes', level=2)
add_para(
    "The primary outcome was the natural logarithm of monthly approved quantity (log_qtd), measured at "
    "the procedure subgroup level. Secondary outcomes included total approved value in constant December "
    "2024 reais and the number of active healthcare establishments by type and legal nature."
)

doc.add_heading('Statistical analysis', level=2)
add_para(
    "The baseline DiD model was specified as:"
)
add_para(
    "log(Y_it) = \u03b1 + \u03b2\u2081\u00b7Treated_i + \u03b2\u2082\u00b7Post_t + "
    "\u03b2\u2083\u00b7(Treated\u00d7Post)_it + \u03b5_it",
    italic=True
)
add_para(
    "where Y_it is the approved quantity for procedure subgroup i in month t, Treated_i is an indicator "
    "for subgroups that received targeted recomposition, Post_t indicates months after the recomposition "
    "event, and \u03b2\u2083 is the DiD estimator of interest. We estimated progressively richer "
    "specifications: (M1a) basic DiD, (M1b) DiD with procedure fixed effects and a linear time trend, "
    "and (M1c) DiD with procedure and time (year-month) fixed effects. All models used "
    "heteroskedasticity-consistent standard errors (HC1) [27]."
)
add_para(
    "For the 2023 recomposition, we estimated a parallel model (M2) comparing the 2023 treatment group "
    "with the control group. For ambulatory production, model M3 replicated the M1c specification using "
    "SIA/SUS data."
)
add_para(
    "To assess regional heterogeneity, we estimated separate DiD models for each of 15 Brazilian states "
    "(M4), examining whether recomposition effects varied by state. Results are presented as a forest plot."
)
add_para(
    "To validate the parallel trends assumption, we estimated an event study model (M5) with semester-level "
    "interaction terms between the treatment indicator and period dummies, using the first semester of 2021 "
    "as the reference period. Additionally, a placebo test (M6) applied a false treatment date (July 2021) "
    "to pre-recomposition data to verify that no spurious effect emerged."
)

doc.add_heading('Ethical considerations', level=2)
add_para(
    "This study used exclusively publicly available, anonymized administrative data. No individual-level "
    "patient data were accessed. In accordance with Brazilian regulations (Resolution CNS 510/2016), "
    "studies using public-domain secondary data do not require ethics committee approval [28]."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Results', level=1)

doc.add_heading('SIGTAP reference values: nominal and real trends', level=2)
add_para(
    "Table 1 presents the evolution of SIGTAP reference values for the 26 procedure subgroups analyzed. "
    "Between January 2018 and September 2025, nominal values increased for all subgroups, but the magnitude "
    "varied substantially by recomposition history. Subgroups in the 2024 treatment group experienced the "
    "largest nominal increases (35.0\u201355.9%), led by tomography and MRI (55.9% each) and followed by "
    "ophthalmologic and cardiovascular surgery (48.5% each). The 2023 treatment group saw increases of "
    "29.4\u201335.8%. Control subgroups received only the 12.5% across-the-board adjustment in 2025."
)
add_para(
    "After deflation by the IPCA (base December 2024 = 100), the picture changed dramatically. "
    "Tomography and MRI were the only subgroups to achieve positive real-value growth (+4.2%), reflecting "
    "their cumulative targeted recompositions in 2019, 2022, and 2024. All other subgroups experienced "
    "real-value erosion: cardiovascular and ophthalmologic surgery lost 0.8% in real terms; oncology lost "
    "9.2%; and control subgroups suffered the largest real losses, ranging from \u221222.6% (consultations, "
    "physiotherapy) to \u221224.8% (anatomical pathology, radiology, endoscopy, and others). Figure 1 "
    "illustrates these nominal-versus-real trajectories for six representative subgroups."
)

# Tabela 2 - Recomposição
add_table_from_csv(TAB_DIR / "tabela2_recomposicao_sigtap.csv",
                   "Table 1. Accumulated nominal and real variation of SIGTAP reference values by procedure "
                   "subgroup, January 2018 to September 2025.")

# Figura 1
add_figure(FIG_DIR / "fig1_sigtap_series.png",
           "Figure 1. Nominal and real (December 2024 R$) SIGTAP reference values for selected procedure "
           "subgroups, Brazil, 2018\u20132025. Dashed vertical lines indicate recomposition events.")

doc.add_heading('Hospital production (SIH/SUS)', level=2)
add_para(
    "Table 2 summarizes hospital production by DiD group and period. All groups experienced sharp declines "
    "during the COVID-19 pandemic (March 2020\u2013June 2021), with mean monthly quantities falling by "
    "approximately 21\u201322% relative to pre-COVID levels. Recovery occurred progressively, with production "
    "exceeding pre-COVID levels by mid-2022 for most subgroups."
)
add_para(
    "In the post-recomposition 2024 period (March 2024\u2013September 2025), the treatment group showed "
    "mean monthly production 30.1% above pre-COVID levels, compared to 14.8% for the control group. "
    "The largest individual increases were observed in cardiovascular surgery (+33.1%), digestive surgery "
    "(+32.0%), and transplantation (+30.0%). Oncology (+24.4%), nephrology (+24.0%), and oncologic surgery "
    "(+25.2%) also showed substantial gains in the 2023 treatment group. Control subgroups showed more "
    "modest increases of 12.2\u201317.5%, consistent with underlying secular trends."
)

add_table_from_csv(TAB_DIR / "tabela3_resumo_did.csv",
                   "Table 2. Summary statistics of hospital production (SIH/SUS) by DiD group and period, "
                   "Brazil, 2018\u20132025.")

add_table_from_csv(TAB_DIR / "tabela4_variacao_pre_pos.csv",
                   "Table 3. Percentage variation in mean monthly hospital production: post-recomposition "
                   "(2024/03\u20132025/09) vs. pre-COVID (2018/01\u20132020/02), by procedure subgroup.")

# Figura 2
add_figure(FIG_DIR / "fig2_sih_did_index.png",
           "Figure 2. Indexed hospital production (January 2018 = 100) for treatment (2024 recomposition) "
           "and control groups. LOESS smoothing applied.")

doc.add_heading('Ambulatory production (SIA/SUS)', level=2)
add_para(
    "Ambulatory production showed similar patterns. COVID-19 disruption was more severe in ambulatory "
    "settings, with quantities falling by up to 55% in March\u2013May 2020 before gradual recovery. "
    "Subgroups receiving targeted recomposition (tomography, MRI, nuclear medicine) showed accelerated "
    "recovery relative to control subgroups, with the divergence becoming more pronounced after the 2024 "
    "recomposition. Figure 3 shows individual subgroup trajectories."
)

add_figure(FIG_DIR / "fig3_sia_subgrupos.png",
           "Figure 3. Monthly ambulatory production (SIA/SUS) by procedure subgroup, Brazil, 2018\u20132025. "
           "Colors indicate DiD group classification.", width=6.0)

doc.add_heading('Difference-in-differences estimates', level=2)
add_para(
    "Table 4 presents the results of the five DiD models. The basic specification (M1a) yielded a DiD "
    "coefficient of 0.120 but was not statistically significant (p = 0.569) due to high residual variance "
    "without fixed effects. After including procedure fixed effects and a linear trend (M1b), the coefficient "
    "increased to 0.161 (p < 0.001), reflecting the correction for compositional differences. The preferred "
    "specification (M1c), with procedure and time fixed effects, estimated a DiD coefficient of 0.120 "
    "(SE = 0.005; t = 24.6; p < 0.001; R\u00b2 = 0.9996), corresponding to an approximately 12.7% increase "
    "in approved hospital production attributable to the 2024 recomposition (exp(0.120) \u2212 1 = 0.127)."
)
add_para(
    "The 2023 recomposition (M2) showed a smaller but significant effect: coefficient = 0.081 (SE = 0.005; "
    "p < 0.001), corresponding to an 8.4% increase. The ambulatory DiD (M3) closely mirrored the hospital "
    "estimate: coefficient = 0.120 (SE = 0.005; p < 0.001), suggesting that recomposition effects were "
    "consistent across care settings."
)

add_table_from_csv(TAB_DIR / "tabela5_modelos_did.csv",
                   "Table 4. Difference-in-differences regression results for the effect of fee recomposition "
                   "on approved production (log quantity). HC1 robust standard errors.")

doc.add_heading('Validation: event study and placebo test', level=2)
add_para(
    "The event study (Figure 4) shows the semester-level DiD coefficients relative to the first semester "
    "of 2021 (reference period). Pre-treatment coefficients fluctuate near zero (apart from the COVID "
    "disruption in 2020), supporting the parallel trends assumption. A visible upward shift occurs after "
    "the 2024 recomposition (semesters 5\u20137), consistent with a treatment effect."
)
add_para(
    "The placebo test (Table 5) applied a false treatment date of July 2021 to the pre-2024 data. "
    "The estimated coefficient was 0.003 (SE = 0.004; p = 0.488), confirming that no spurious effect "
    "existed prior to the actual recomposition. This result strengthens the causal interpretation of "
    "the DiD estimates."
)

add_figure(FIG_DIR / "fig7_event_study.png",
           "Figure 4. Event study: semester-level DiD coefficients for hospital production. "
           "Reference period: 2021-S1. Shaded area indicates 95% confidence intervals.")

add_table_from_csv(TAB_DIR / "tabela7_placebo.csv",
                   "Table 5. Placebo test results: false treatment date (July 2021) applied to "
                   "pre-recomposition data.")

doc.add_heading('Regional heterogeneity', level=2)
add_para(
    "Figure 5 presents the state-level DiD estimates for the 2024 recomposition. All 15 states showed "
    "positive and statistically significant effects (p < 0.05). The magnitude ranged from 9.1% in "
    "Goi\u00e1s (coefficient = 0.091; 95% CI: 0.061\u20130.120) to 13.1% in Maranh\u00e3o "
    "(coefficient = 0.131; 95% CI: 0.103\u20130.159). Notably, states with historically lower per-capita "
    "supply (Maranh\u00e3o, Pernambuco, Bahia, Pará) tended to show larger effects, though the "
    "differences between states were not statistically significant given overlapping confidence intervals."
)

add_figure(FIG_DIR / "fig8_forest_uf.png",
           "Figure 5. Forest plot: state-level DiD estimates for the 2024 recomposition effect on "
           "hospital production. HC1 robust standard errors, 95% CI.")

add_table_from_csv(TAB_DIR / "tabela6_efeitos_uf.csv",
                   "Table 6. State-level DiD estimates for the 2024 recomposition effect on hospital "
                   "production (log quantity).")

doc.add_heading('Healthcare establishment network (CNES)', level=2)
add_para(
    "The CNES analysis revealed a modest but consistent expansion of the healthcare establishment network "
    "over the study period. Between 2018 and 2025, the total number of active specialized care facilities "
    "(clinics, diagnostic centers, and hospitals) grew by approximately 11\u201315%. Post-recomposition, "
    "specialty clinics and diagnostic/therapeutic support services (SADT) showed an additional 4\u20136% "
    "increase relative to pre-2024 trends (Figure 6). The expansion was concentrated in private for-profit "
    "and philanthropic facilities, while public establishments remained relatively stable."
)

add_figure(FIG_DIR / "fig4_cnes_evolucao.png",
           "Figure 6. Evolution of active healthcare establishments by type and legal nature (CNES), "
           "Brazil, 2018\u20132025 (quarterly).")

doc.add_heading('Total production value', level=2)
add_para(
    "Figure 7 shows the total real value of approved production across SIH and SIA. The combined real "
    "production value declined sharply during the pandemic but showed sustained recovery from late 2020, "
    "surpassing pre-COVID levels by mid-2022. The acceleration after the 2024 recomposition is visible, "
    "though it reflects both volume and price effects. Hospital production (SIH) accounts for approximately "
    "70\u201375% of total specialized care spending."
)

add_figure(FIG_DIR / "fig6_valor_total.png",
           "Figure 7. Total approved production value in constant December 2024 R$ (billions), "
           "SIH/SUS and SIA/SUS, Brazil, 2018\u20132025.")

# Figura defasagem real
add_figure(FIG_DIR / "fig5_defasagem_real.png",
           "Figure 8. Accumulated real-value variation of SIGTAP reference values relative to January 2018, "
           "selected subgroups. IPCA-deflated.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Discussion', level=1)

add_para(
    "This study provides the first quasi-experimental evidence that fee schedule recomposition in the "
    "Brazilian SUS is associated with meaningful increases in specialized care production. The 2024 "
    "recomposition was associated with a 12\u201313% increase in both hospital and ambulatory approved "
    "production, while the 2023 recomposition showed an 8% increase. These effects were consistent across "
    "care settings, robust to different model specifications, and validated by both an event study and "
    "a placebo test."
)

doc.add_heading('Comparison with international literature', level=2)
add_para(
    "Our findings are broadly consistent with the international literature on the price-quantity relationship "
    "in healthcare. Clemens and Gottlieb (2014) found that a 2% Medicare fee increase led to approximately "
    "3% more healthcare spending in the United States, implying an elasticity of roughly 1.5 [13]. Our "
    "estimates suggest a somewhat lower elasticity: the 20% fee increase in the 2024 recomposition was "
    "associated with a 12\u201313% production increase, implying an elasticity of approximately 0.6. This "
    "difference may reflect the specific institutional context of the SUS, where provider participation is "
    "shaped by factors beyond price\u2014including contractual obligations, regulatory requirements, and "
    "the mixed public-private nature of the Brazilian health system [2, 20]."
)
add_para(
    "Decker (2009) showed that Medicaid fee increases from 1996 to 2004 were associated with increased "
    "physician willingness to accept new Medicaid patients [12]. While our data do not directly capture "
    "provider willingness, the CNES analysis suggests that the establishment network expanded modestly "
    "post-recomposition, particularly among private and philanthropic facilities\u2014consistent with a "
    "positive supply response to improved reimbursement."
)
add_para(
    "In the European context, Januleviciute et al. (2016) documented that the introduction of "
    "activity-based financing in Norwegian hospitals altered production volumes across diagnostic categories "
    "[15]. Our event study similarly captures a post-treatment divergence that is consistent with a causal "
    "effect of fee changes on production."
)

doc.add_heading('Implications for Brazilian health policy', level=2)
add_para(
    "The central contribution of this study is to shift the debate about the Tabela SUS from the "
    "exclusively rhetorical plane to the empirical measurement plane. Our results demonstrate that "
    "administered prices matter: when the SUS increases reimbursement for specific procedure categories, "
    "approved production responds positively. This finding has several policy implications."
)
add_para(
    "First, it supports the use of targeted fee recomposition as a policy instrument for expanding access "
    "to specific types of specialized care, particularly in settings where the gap between administered "
    "values and production costs has widened most. The larger effects observed in states with historically "
    "lower supply (Maranh\u00e3o, Pernambuco, Par\u00e1) suggest that recomposition may be particularly "
    "effective in underserved areas where provider margins have been most compressed [8, 29]."
)
add_para(
    "Second, the finding that real values continued to decline for most subgroups despite nominal increases "
    "underscores the inadequacy of sporadic, one-time adjustments. A systematic mechanism for periodic "
    "inflation-indexed recomposition\u2014analogous to automatic adjustment clauses in regulated industries"
    "\u2014would better preserve the real value of the fee schedule over time [5, 30]."
)
add_para(
    "Third, the modest magnitude of the CNES network expansion suggests that fee recomposition alone is "
    "insufficient to substantially restructure the provider network. Complementary strategies are needed, "
    "including strengthened contracting mechanisms, regional health planning, investment in installed "
    "capacity, and queue management systems [2, 7, 31]."
)

doc.add_heading('Literature gaps addressed', level=2)
add_para(
    "This study fills several gaps identified in our review of the literature:"
)
add_para(
    "1. Measurement of real-value erosion: We constructed the first comprehensive deflated time series "
    "of SIGTAP reference values from 2018 to 2025, demonstrating that real values declined by up to 25% "
    "for subgroups without targeted recomposition. Previous studies cited fee inadequacy qualitatively "
    "but lacked systematic quantification [5, 6, 18]."
)
add_para(
    "2. Causal identification: By applying a DiD design with procedure and time fixed effects, we moved "
    "beyond the correlational evidence that characterized prior Brazilian studies [6, 19]. The event study "
    "and placebo test provide additional support for causal interpretation."
)
add_para(
    "3. Regional heterogeneity: Our state-level analysis reveals that recomposition effects are not "
    "uniform across Brazil, with suggestive evidence that effects are larger in historically underserved "
    "states\u2014a finding with direct relevance for equity-oriented resource allocation [29, 32]."
)
add_para(
    "4. Multi-database integration: By linking SIGTAP values with SIH/SUS production, SIA/SUS production, "
    "and CNES establishment data, we provide a more comprehensive picture than studies relying on a single "
    "data source [22, 23, 24]."
)
add_para(
    "5. LMIC evidence on administered prices: This study contributes to the sparse international literature "
    "on the price-production nexus in low- and middle-income country health systems, where most evidence "
    "originates from US Medicare/Medicaid or European single-payer contexts [16, 17]."
)

doc.add_heading('Limitations', level=2)
add_para(
    "Several limitations should be acknowledged. First, while the DiD design provides stronger causal "
    "identification than cross-sectional approaches, it relies on the parallel trends assumption. Although "
    "our event study and placebo test support this assumption, unmeasured confounders correlated with both "
    "recomposition timing and production trends cannot be entirely ruled out."
)
add_para(
    "Second, the study analyzes approved production (quantities approved for reimbursement), not actual "
    "healthcare delivery. Processing delays, billing practices, and authorization patterns may introduce "
    "noise. We excluded the final months of the series to mitigate processing lag, but residual bias "
    "cannot be excluded."
)
add_para(
    "Third, the CNES analysis is limited to quarterly snapshots of establishment counts and does not "
    "capture within-facility changes in capacity, equipment, or staffing. A finer-grained analysis of "
    "installed capacity\u2014including beds, equipment units, and professional hours\u2014would provide "
    "a more complete picture of supply-side responses."
)
add_para(
    "Fourth, this study does not analyze the quality or appropriateness of the additional production "
    "generated by recomposition. An increase in approved quantities could reflect both genuine expansion "
    "of beneficial care and potential supply-induced demand or coding changes [33]."
)
add_para(
    "Fifth, the analysis is conducted at the procedure subgroup level, aggregating heterogeneous procedures "
    "within each subgroup. Subgroup-level effects may mask important variation at the individual procedure "
    "code level."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# CONCLUSIONS
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Conclusions', level=1)
add_para(
    "This study provides robust empirical evidence that fee schedule recomposition in the SUS is "
    "associated with significant increases in specialized care production. The 2024 recomposition was "
    "associated with a 12\u201313% increase in both hospital and ambulatory approved production, with "
    "effects detectable across all states analyzed. The 2023 recomposition showed a smaller but "
    "significant 8% effect."
)
add_para(
    "These findings support the hypothesis that administered prices matter for healthcare supply in "
    "publicly financed systems. However, the persistence of real-value erosion in most procedure "
    "categories\u2014even after nominal recomposition\u2014highlights the need for systematic, "
    "inflation-indexed adjustment mechanisms. Recomposition is a necessary but insufficient condition "
    "for expanding access to specialized care; complementary investments in regional planning, provider "
    "contracting, installed capacity, and queue management remain essential."
)
add_para(
    "Future research should examine the quality and appropriateness of production changes following "
    "recomposition, explore the mechanisms through which fee changes translate into provider behavior, "
    "and evaluate the cost-effectiveness of targeted recomposition as a policy instrument for expanding "
    "access to specialized care in the SUS."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DECLARATIONS
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Declarations', level=1)

doc.add_heading('Ethics approval and consent to participate', level=2)
add_para(
    "Not applicable. This study used exclusively publicly available, anonymized secondary data. "
    "In accordance with Brazilian National Health Council Resolution 510/2016, ethics committee "
    "approval is not required for research using public-domain data."
)

doc.add_heading('Consent for publication', level=2)
add_para("Not applicable.")

doc.add_heading('Availability of data and materials', level=2)
add_para(
    "All data used in this study are publicly available through the Brazilian Ministry of Health\u2019s "
    "DATASUS platform (https://datasus.saude.gov.br). The analytical code and processed datasets are "
    "available in the supplementary repository."
)

doc.add_heading('Competing interests', level=2)
add_para("The authors declare that they have no competing interests.")

doc.add_heading('Funding', level=2)
add_para("[To be completed by authors]")

doc.add_heading('Authors\u2019 contributions', level=2)
add_para("[To be completed by authors]")

doc.add_heading('Acknowledgements', level=2)
add_para("[To be completed by authors]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════

doc.add_heading('References', level=1)

references = [
    "1. Paim J, Travassos C, Almeida C, Bahia L, Macinko J. The Brazilian health system: history, advances, and challenges. Lancet. 2011;377(9779):1778\u201397.",
    "2. Gragnolati M, Lindelow M, Couttolenc B. Twenty years of health system reform in Brazil: an assessment of the Sistema Único de Saúde. Washington, DC: World Bank; 2013.",
    "3. Brasil. Ministério da Saúde. Sistema de Gerenciamento da Tabela de Procedimentos, Medicamentos e OPM do SUS (SIGTAP). Brasília: Ministério da Saúde; 2025. Available from: http://sigtap.datasus.gov.br",
    "4. Santos L, Andrade LOM. SUS: o espaço da gestão inovada e dos consensos interfederativos. Campinas: Saberes; 2011.",
    "5. Conselho Nacional de Secretários de Saúde (CONASS). A Tabela SUS e a remuneração dos procedimentos: nota técnica. Brasília: CONASS; 2023.",
    "6. Portela MC, Proença JD, Lima SML, Barbosa PR, Vasconcellos MM, Ugá MAD, et al. Estudo sobre os hospitais filantrópicos no Brasil. Rev Saúde Pública. 2018;52:34.",
    "7. Solla JJSP, Chioro A. Atenção ambulatorial especializada. In: Giovanella L, Escorel S, Lobato LVC, Noronha JC, Carvalho AI, editors. Políticas e sistema de saúde no Brasil. 2nd ed. Rio de Janeiro: Fiocruz; 2012. p. 547\u201376.",
    "8. Confederação Nacional de Saúde (CNSaúde). O impacto da defasagem da Tabela SUS na sustentabilidade dos hospitais. Brasília: CNSaúde; 2022.",
    "9. McGuire TG. Physician agency and payment mechanisms. In: Culyer AJ, Newhouse JP, editors. Handbook of Health Economics. Vol. 1A. Amsterdam: Elsevier; 2000. p. 461\u2013536.",
    "10. Ellis RP, McGuire TG. Provider behavior under prospective reimbursement: cost sharing and supply. J Health Econ. 1986;5(2):129\u201351.",
    "11. Rice T, Labelle RJ. Do physicians induce demand for medical services? J Health Polit Policy Law. 1989;14(3):587\u2013600.",
    "12. Decker SL. Changes in Medicaid physician fees and patterns of ambulatory care. Inquiry. 2009;46(3):291\u2013304.",
    "13. Clemens J, Gottlieb JD. Do physicians' financial incentives affect medical treatment and patient health? Am Econ Rev. 2014;104(4):1320\u201349.",
    "14. Zuckerman S, Skopec L, Epstein M. Medicaid physician fees after the ACA primary care fee bump. Urban Institute; 2017.",
    "15. Januleviciute J, Askildsen JE, Kaarboe O, Siciliani L, Sutton M. How do hospitals respond to price changes? Evidence from Norway. Health Econ. 2016;25(5):620\u201336.",
    "16. Wagstaff A, Eozenou PHV, Smitz MF. Out-of-pocket expenditures on health: a global stocktake. World Bank Res Obs. 2020;35(2):123\u201357.",
    "17. Mills A, Ataguba JE, Akazili J, Borghi J, Garshong B, Makawia S, et al. Equity in financing and use of health care in Ghana, South Africa, and Tanzania: implications for paths to universal coverage. Lancet. 2012;380(9837):126\u201333.",
    "18. Santos MAB, Ugá MAD. Reformas organizacionais no setor saúde: os estímulos e os incentivos da política de remuneração dos serviços hospitalares no Brasil (1984\u20132004). Saúde em Debate. 2007;31(75\u201377):30\u201343.",
    "19. Machado JP, Martins M, Leite IC, Silva GS, Oliveira RB. O perfil dos hospitais brasileiros que atendem ao SUS. Cad Saúde Pública. 2014;30(Suppl):S72\u2013S86.",
    "20. La Forgia GM, Couttolenc BF. Hospital performance in Brazil: the search for excellence. Washington, DC: World Bank; 2008.",
    "21. Couttolenc BF, Gragnolati M. Performance of Brazil\u2019s public health system. In: Twenty years of health system reform in Brazil. Washington, DC: World Bank; 2011.",
    "22. Brasil. Ministério da Saúde. Sistema de Informações Hospitalares do SUS (SIH/SUS). Brasília: DATASUS; 2025. Available from: https://datasus.saude.gov.br",
    "23. Brasil. Ministério da Saúde. Sistema de Informações Ambulatoriais do SUS (SIA/SUS). Brasília: DATASUS; 2025. Available from: https://datasus.saude.gov.br",
    "24. Brasil. Ministério da Saúde. Cadastro Nacional de Estabelecimentos de Saúde (CNES). Brasília: DATASUS; 2025. Available from: https://cnes.datasus.gov.br",
    "25. Instituto Brasileiro de Geografia e Estatística (IBGE). Índice Nacional de Preços ao Consumidor Amplo \u2013 IPCA. Rio de Janeiro: IBGE; 2025.",
    "26. Brasil. Ministério da Saúde. Portaria GM/MS nº 6.465, de 30 de dezembro de 2024. Dispõe sobre a recomposição dos valores da Tabela de Procedimentos, Medicamentos e OPM do SUS. Diário Oficial da União; 2024.",
    "27. White H. A heteroskedasticity-consistent covariance matrix estimator and a direct test for heteroskedasticity. Econometrica. 1980;48(4):817\u201338.",
    "28. Brasil. Conselho Nacional de Saúde. Resolução nº 510, de 7 de abril de 2016. Diário Oficial da União; 2016.",
    "29. Viacava F, Oliveira RAD, Carvalho CC, Laguardia J, Bellido JG. SUS: oferta, acesso e utilização de serviços de saúde nos últimos 30 anos. Ciênc Saúde Coletiva. 2018;23(6):1751\u201362.",
    "30. Vieira FS, Benevides RPS. Os impactos do novo regime fiscal para o financiamento do Sistema Único de Saúde e para a efetivação do direito à saúde no Brasil. Nota Técnica IPEA; 2016.",
    "31. Giovanella L, Mendonça MHM, Buss PM, Fleury S, Gadelha CAG, Galvão LAC, et al. De Alma-Ata a Astana: atenção primária à saúde e sistemas universais de saúde. Cad Saúde Pública. 2019;35(3):e00066019.",
    "32. Albuquerque MV, Viana ALD, Lima LD, Ferreira MP, Fusaro ER, Iozzi FL. Desigualdades regionais na saúde: mudanças observadas no Brasil de 2000 a 2016. Ciênc Saúde Coletiva. 2017;22(4):1015\u201328.",
    "33. Gruber J, Owings M. Physician financial incentives and cesarean section delivery. RAND J Econ. 1996;27(1):99\u2013123.",
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)

# ── Salvar ────────────────────────────────────────────────────────
output_path = DOC_DIR / "manuscrito_tabela_sus_bmc.docx"
doc.save(str(output_path))
print(f"Manuscrito salvo em: {output_path}")
